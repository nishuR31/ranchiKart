import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideV w:val="none"/>\n'
        f'  <w:left w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_header_footer(doc):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("RanchiKart — SIP Project Report | The ICFAI University Jharkhand")
        hrun.font.name = "Calibri"
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("Department of Computer Science & IT — The ICFAI University Jharkhand, Ranchi")
        frun.font.name = "Calibri"
        frun.font.size = Pt(9)
        frun.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

def style_heading(p, text, level, font_name="Calibri"):
    p.text = text
    run = p.runs[0]
    run.font.name = font_name
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78) # Deep Navy
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        run.font.size = Pt(11.5)
        run.font.color.rgb = RGBColor(0x59, 0x59, 0x59) # Slate Gray
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)

def add_styled_paragraph(doc, text, space_after=6, line_spacing=1.15, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.name = "Calibri"
        r_pre.font.size = Pt(11)
        r_pre.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    return p

def add_styled_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.name = "Calibri"
        r_pre.font.size = Pt(11)
        r_pre.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    return p

def create_styled_table(doc, headers, rows_data, col_widths=None):
    table = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # Header Row
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        set_cell_background(hdr_cells[i], "1F4E78") # Deep Navy
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.name = "Calibri"
            r.font.size = Pt(10)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data Rows
    for r_idx, row_values in enumerate(rows_data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F9FAFB" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_values):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=150, right=150)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = "Calibri"
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(6)
    return table

def build_sip_report():
    doc = Document()
    add_header_footer(doc)

    # ---------------------------------------------------------
    # 1.0 COVER PAGE
    # ---------------------------------------------------------
    p_univ = doc.add_paragraph()
    p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_univ.paragraph_format.space_before = Pt(36)
    r_univ = p_univ.add_run("THE ICFAI UNIVERSITY JHARKHAND, RANCHI")
    r_univ.bold = True
    r_univ.font.name = "Arial"
    r_univ.font.size = Pt(18)
    r_univ.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(36)
    r_sub = p_sub.add_run("Faculty of Science & Technology | Department of Computer Science & IT")
    r_sub.font.name = "Calibri"
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(12)
    r_title = p_title.add_run("SUMMER INTERNSHIP PROGRAM (SIP) PROJECT REPORT\nON")
    r_title.bold = True
    r_title.font.name = "Arial"
    r_title.font.size = Pt(14)
    r_title.font.color.rgb = RGBColor(0x26, 0x26, 0x26)

    p_proj = doc.add_paragraph()
    p_proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_proj.paragraph_format.space_after = Pt(48)
    r_proj = p_proj.add_run("RanchiKart (BasketByte)\nFull-Stack Hyperlocal E-Commerce Web Portal")
    r_proj.bold = True
    r_proj.font.name = "Arial"
    r_proj.font.size = Pt(20)
    r_proj.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    # Details table on cover page
    cover_data = [
        ["Submitted By (Group Members)", "Nishan · Aditya · Suraj · Tushar · Adarsh"],
        ["Program", "Bachelor of Technology (B.Tech) / MCA / BCA"],
        ["Faculty Guide", "Prof. Abhay Kumar Sinha (Faculty of Science & Technology)"],
        ["Company / Industry Guide", "Prof. Abhay Kumar Sinha / Technical Lead"],
        ["Academic Year", "2024 – 2026"],
        ["Submission Date", "July 2026"]
    ]
    create_styled_table(doc, ["Attribute", "Details"], cover_data, [2.5, 4.0])
    doc.add_page_break()

    # ---------------------------------------------------------
    # 2.0 TABLE OF CONTENTS
    # ---------------------------------------------------------
    p_toc_head = doc.add_paragraph()
    style_heading(p_toc_head, "2.0 TABLE OF CONTENTS", 1)
    
    toc_items = [
        ("1.0 Cover Page", "1"),
        ("2.0 Table of Contents", "2"),
        ("3.0 Acknowledgment", "3"),
        ("4.0 Introduction and Objective", "4"),
        ("5.0 Brief Introduction of Company (Development Environment)", "5"),
        ("6.0 Brief Introduction of Client Company", "5"),
        ("7.0 Project Detail", "6"),
        ("    7.1 Application Overview & Key Metadata", "6"),
        ("    7.2 Role & Contribution (Group Responsibility Distribution)", "6"),
        ("    7.3 Study of Existing System", "7"),
        ("    7.4 Details of Shortcomings of Existing System", "7"),
        ("    7.5 Details of Fresh Requirements from New System", "7"),
        ("    7.6 The Proposed System (Architecture & Requirements Match)", "8"),
        ("8.0 Feasibility Study (Cost, Benefit & Risk Analysis)", "9"),
        ("9.0 System Details (Details of Application Developed)", "11"),
        ("    9.1 Overview & Scope", "11"),
        ("    9.2 Application Modules", "11"),
        ("    9.3 Tools & Technology (Software & Hardware)", "12"),
        ("    9.4 System Design", "13"),
        ("        9.4.1 Data Flow Diagrams (DFD)", "13"),
        ("        9.4.2 Entity-Relationship (ER) Diagram", "14"),
        ("        9.4.3 Database Design (13 Tables, Storage Math, Procedures, Triggers, Roles)", "14"),
        ("        9.4.4 Form Design (Validation, Calculations, Data Fetching, Navigation)", "18"),
        ("        9.4.5 Data Processing (Background Jobs & Cleanup)", "21"),
        ("        9.4.6 Networking Details & Protocols", "21"),
        ("        9.4.7 Security Features & Vulnerability Mitigation", "22"),
        ("10.0 System Testing (Strategy, Test Cases & CI Report)", "24"),
        ("11.0 Scope of Our System", "26"),
        ("12.0 Future Implementation & Roadmap", "26"),
        ("13.0 Bibliography & References", "27"),
        ("Appendix A — System Flowchart", "28"),
        ("Appendix B — Project Responsibility Document (Full)", "29")
    ]
    
    for item, pg in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(item)
        r1.font.name = "Calibri"
        r1.font.size = Pt(11)
        r1.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
        # Dots leader
        dots_count = max(5, 75 - len(item))
        r_dots = p.add_run(" " + "." * dots_count + " ")
        r_dots.font.name = "Calibri"
        r_dots.font.size = Pt(10)
        r_dots.font.color.rgb = RGBColor(0xA6, 0xA6, 0xA6)
        
        r2 = p.add_run(pg)
        r2.bold = True
        r2.font.name = "Calibri"
        r2.font.size = Pt(11)

    doc.add_page_break()

    # ---------------------------------------------------------
    # 3.0 ACKNOWLEDGMENT
    # ---------------------------------------------------------
    p_ack_head = doc.add_paragraph()
    style_heading(p_ack_head, "3.0 ACKNOWLEDGMENT", 1)
    
    add_styled_paragraph(doc, "We express our sincere gratitude to The ICFAI University Jharkhand, Ranchi, and the Faculty of Science & Technology for providing us with the opportunity to undertake this Summer Internship Program (SIP) project titled RanchiKart (BasketByte).")
    add_styled_paragraph(doc, "We extend our deepest gratitude to our Faculty Guide, Prof. Abhay Kumar Sinha, for his invaluable supervision, continuous encouragement, constructive guidance, and meticulous review throughout the lifecycle of this project. His insights significantly refined our engineering approach, system architecture, and technical documentation.")
    add_styled_paragraph(doc, "We also thank our team members — Nishan, Aditya, Suraj, Tushar, and Adarsh — whose seamless collaboration across backend engineering, frontend UI/UX development, software testing, system design, and project coordination made the successful completion and live deployment of RanchiKart possible.")

    # ---------------------------------------------------------
    # 4.0 INTRODUCTION AND OBJECTIVE
    # ---------------------------------------------------------
    p_intro_head = doc.add_paragraph()
    style_heading(p_intro_head, "4.0 INTRODUCTION AND OBJECTIVE", 1)
    
    add_styled_paragraph(doc, "RanchiKart (internally codenamed BasketByte) is a full-stack, B2C e-commerce web platform engineered specifically for the regional market of Ranchi, Jharkhand. The primary motivation behind RanchiKart is bridging the digital divide for local MSME merchants, stationery providers, stamp makers, and specialty goods sellers who lack access to modern, high-performance digital commerce infrastructure.")
    
    style_heading(doc.add_paragraph(), "Core Objectives of the System", 2)
    add_styled_bullet(doc, "Deliver a unified, secure, high-speed shopping experience for local buyers with sub-second page loads and mobile-responsive rendering.", "1. Hyperlocal Accessibility: ")
    add_styled_bullet(doc, "Provide specialized order item customisation (custom width, height, and custom text engraving) required for made-to-order products like customized rubber stamps and acrylic name boards.", "2. Made-to-Order Customisation: ")
    add_styled_bullet(doc, "Support five modern authentication mechanisms (Passwords with Argon2id, Google OAuth 2.0, Passwordless Magic Links, TOTP 2FA, and WebAuthn / FIDO2 Passkeys) unified under a stateless JWT session architecture.", "3. Multi-Method Security: ")
    add_styled_bullet(doc, "Integrate Razorpay payment gateway with cryptographic HMAC-SHA256 signature verification and strict payment state transitions.", "4. Financial & Payment Integrity: ")
    add_styled_bullet(doc, "Equip store managers and system administrators with complete operational control over products, orders, categories, coupons, user access, and system audit logs.", "5. Comprehensive Back-Office Admin: ")

    # ---------------------------------------------------------
    # 5.0 & 6.0 COMPANY PROFILES
    # ---------------------------------------------------------
    style_heading(doc.add_paragraph(), "5.0 BRIEF INTRODUCTION OF THE COMPANY (DEVELOPMENT ENVIRONMENT)", 1)
    add_styled_paragraph(doc, "The project was developed as part of the Summer Internship Program (SIP) initiative under the aegis of the Department of Computer Science & IT, Faculty of Science & Technology, The ICFAI University Jharkhand, Ranchi. The software engineering environment modeled production-grade industry standard workflows, employing Git-based version control, continuous integration (CI) via GitHub Actions, containerized backend services using Docker, and cloud deployments across Vercel (frontend), Render (API backend), and NeonDB (PostgreSQL Serverless DB).")

    style_heading(doc.add_paragraph(), "6.0 BRIEF INTRODUCTION OF THE CLIENT COMPANY", 1)
    add_styled_paragraph(doc, "The client persona for RanchiKart comprises small-to-medium enterprise (SME) local merchants, independent retail shopkeepers, stationery vendors, and customized product manufacturers situated in Ranchi, Jharkhand. These businesses operate without dedicated IT departments and require a zero-maintenance, zero-commission-barrier storefront that handles cataloging, online payment collection, and customer order management effortlessly.")

    # ---------------------------------------------------------
    # 7.0 PROJECT DETAIL
    # ---------------------------------------------------------
    style_heading(doc.add_paragraph(), "7.0 PROJECT DETAIL", 1)
    
    style_heading(doc.add_paragraph(), "7.1 Application Overview & Key Metadata", 2)
    meta_table_data = [
        ["Type of Application", "Full-Stack B2C E-Commerce Web Application (Decoupled SPA + REST API Backend)"],
        ["Platform — Frontend", "React 18.3 SPA (Vite 5 build tool), Zustand 5 state management, Axios HTTP client"],
        ["Platform — Backend", "Node.js-compatible Bun runtime, Fastify 5 web framework, TypeScript 5.8"],
        ["Platform — Database", "PostgreSQL 17 via Prisma ORM 6.9 (NeonDB Serverless / Prisma Accelerate)"],
        ["Platform — Cache & Store", "Redis 7 (Session storage, rate-limiting, token blacklisting via Redis Cloud)"],
        ["Development Duration", "May 2026 – July 2026 (10 Weeks Total)"],
        ["Total Man-Hours Invested", "Approx. 600 Man-Hours (5 Team Members × 10 Weeks × ~12 Hours/Week)"],
        ["Faculty Guide", "Prof. Abhay Kumar Sinha (Department of Computer Science & IT)"],
        ["Industry Guide", "Prof. Abhay Kumar Sinha / Technical Lead"]
    ]
    create_styled_table(doc, ["Attribute / Parameter", "Project Specification"], meta_table_data, [2.5, 4.0])

    style_heading(doc.add_paragraph(), "7.2 Your Role and Contribution (Group Responsibility Distribution)", 2)
    add_styled_paragraph(doc, "RanchiKart was executed by a collaborative five-member engineering team with clear separation of responsibilities across the full software engineering lifecycle:")
    
    team_data = [
        ["Nishan", "Backend Development · DevOps & Deployment · Database Schema Architecture · Security & Auth Engineering"],
        ["Aditya", "Requirement Analysis · UI/UX Design · Frontend Engineering (React/Vite) · State Management · Performance Optimization"],
        ["Suraj", "Quality Assurance & Testing Lead · Functional, UI, Integration, Regression, Responsiveness & E2E Testing"],
        ["Tushar", "Requirement Analysis · System Design & DFD/ER Modeling · Documentation Lead · API Swagger Spec"],
        ["Adarsh", "Project Coordination · Product Management · Content & Category Cataloging · User Acceptance Testing (UAT)"]
    ]
    create_styled_table(doc, ["Team Member", "Primary Responsibility Areas & Technical Contributions"], team_data, [1.8, 4.7])

    style_heading(doc.add_paragraph(), "7.3 Study of Existing System", 2)
    add_styled_paragraph(doc, "Prior to RanchiKart, local merchants in Ranchi relied almost exclusively on physical walk-in sales or informal messaging apps (e.g., WhatsApp). Customers seeking custom products like engraved rubber stamps, brass/acrylic nameplates, and specialized stationery had to physically visit markets, negotiate prices without standardized cataloging, and submit manual text specifications.")

    style_heading(doc.add_paragraph(), "7.4 Shortcomings of the Existing System", 2)
    add_styled_bullet(doc, "Absence of a localized, Ranchi-focused digital marketplace for stationery and specialty merchants.", "1. No Local Digital Platform: ")
    add_styled_bullet(doc, "National e-commerce platforms (Amazon, Flipkart) do not support dynamic width, height, and custom text inputs per order item.", "2. Lack of Customization Support: ")
    add_styled_bullet(doc, "Exorbitant seller commission rates (15-30%) on major platforms are economically unviable for small local shopkeepers.", "3. High Marketplace Fees: ")
    add_styled_bullet(doc, "Informal WhatsApp orders result in lost tracking, payment disputes, and unorganized customer communication.", "4. Unstructured Ordering: ")
    add_styled_bullet(doc, "Small vendors lack the technical capability to implement secure 2FA, WebAuthn passkeys, or audited admin systems.", "5. Inadequate Security: ")

    style_heading(doc.add_paragraph(), "7.5 Fresh Requirements from the New System", 2)
    add_styled_bullet(doc, "Multi-method authentication supporting Password, Google OAuth, Magic Links, TOTP 2FA, and Passkeys.", "1. Modern Security: ")
    add_styled_bullet(doc, "Dynamic custom width, custom height, and custom text fields stored directly on order line items.", "2. Line-Item Customization: ")
    add_styled_bullet(doc, "Integrated Razorpay payment gateway with cryptographic webhooks and automated status reconciliation.", "3. Payment Integration: ")
    add_styled_bullet(doc, "Centralized Admin Dashboard featuring AdminLog audit tracking, category management, coupon engine, and user control.", "4. Store Back-Office: ")
    add_styled_bullet(doc, "Automated CI/CD pipeline ensuring zero regression releases through fastify.inject integration testing.", "5. Continuous Integration: ")

    style_heading(doc.add_paragraph(), "7.6 The Proposed System", 2)
    add_styled_paragraph(doc, "RanchiKart is architected as a decoupled SPA + REST API platform following the MVCS (Model → Validation → Controller → Service → Repository) design pattern. Request validation is strictly enforced at the Fastify controller level using Zod schemas, while business rules (e.g., coupon eligibility, discount calculations, inventory deduction) reside inside modular Service classes.")
    
    style_heading(doc.add_paragraph(), "How RanchiKart Meets System Requirements", 3)
    req_match_data = [
        ["Secure Multi-Method Login", "Implemented 5 unified auth flows (Password Argon2id, OAuth2, Magic Link, TOTP, Passkeys) under 1 JWT model."],
        ["Custom Product Sizing", "Added customWidthMm, customHeightMm, and customText attributes directly onto OrderItem model."],
        ["Payment Safety", "Razorpay integration with HMAC-SHA256 verification and atomic Prisma database updates."],
        ["Admin Accountability", "AdminLog model recording every administrative action with actor, target entity, and IP/metadata."],
        ["High Performance", "Bun runtime + Fastify framework delivering <50ms API response latency with Redis caching."]
    ]
    create_styled_table(doc, ["Requirement", "How RanchiKart Fulfills It"], req_match_data, [2.2, 4.3])

    # ---------------------------------------------------------
    # 8.0 FEASIBILITY STUDY
    # ---------------------------------------------------------
    style_heading(doc.add_paragraph(), "8.0 FEASIBILITY STUDY", 1)
    
    style_heading(doc.add_paragraph(), "8.1 Cost Estimation of the System", 2)
    cost_data = [
        ["Frontend Hosting (Vercel)", "Global CDN Edge Hosting", "₹0.00 / month (Free Tier)"],
        ["Backend Hosting (Render)", "Web Service Instance (Bun/Node)", "₹0.00 / month (Free Starter)"],
        ["PostgreSQL DB (NeonDB)", "Serverless Postgres + Accelerate", "₹0.00 / month (Free Tier 0.5 GiB)"],
        ["Redis Cache (Redis Cloud)", "In-memory Key-Value Cache", "₹0.00 / month (30 MB Free)"],
        ["Domain Name (.in / .com)", "Annual DNS Registration", "₹800.00 / year"],
        ["Total Operational Overhead", "Production Hosting Stack", "₹800.00 / year (Near Zero)"]
    ]
    create_styled_table(doc, ["Cost Component", "Service Description", "Monthly / Annual Cost"], cost_data, [2.2, 2.5, 1.8])

    style_heading(doc.add_paragraph(), "8.2 Cost-Benefit Analysis", 2)
    cb_data = [
        ["Development Cost (Student Team)", "600 total man-hours invested across 5 team members with zero capital expenditure."],
        ["Infrastructure Overhead", "Leveraged managed cloud free-tiers resulting in zero fixed monthly server maintenance costs."],
        ["Merchant Value Generated", "Eliminated 15-30% platform commissions, allowing local Ranchi shopkeepers 100% margin retention."],
        ["ROI & Payback Period", "Immediate positive return upon first transaction due to negligible operating overhead."]
    ]
    create_styled_table(doc, ["Dimension", "Cost-Benefit Synthesis"], cb_data, [2.0, 4.5])

    style_heading(doc.add_paragraph(), "8.3 Risk Analysis & Mitigation Matrix", 2)
    risk_data = [
        ["Free-tier server cold-starts on Render", "High", "Medium", "Implemented background keep-alive cron pings to keep Fastify warm."],
        ["ImgBB third-party upload downtime", "Medium", "Medium", "Abstracted upload service supporting direct image URL fallback."],
        ["Payment webhook delivery failure", "Low", "High", "Added client-side verification fallback API (/payments/verify) post-checkout."],
        ["Database connection exhaustion", "Medium", "High", "Utilized Prisma Accelerate connection pooling over NeonDB serverless."]
    ]
    create_styled_table(doc, ["Identified Risk Scenario", "Likelihood", "Impact", "Mitigation Strategy"], risk_data, [2.2, 1.0, 1.0, 2.3])

    # ---------------------------------------------------------
    # 9.0 SYSTEM DETAILS
    # ---------------------------------------------------------
    style_heading(doc.add_paragraph(), "9.0 SYSTEM DETAILS (DETAILS OF APPLICATION DEVELOPED)", 1)
    
    style_heading(doc.add_paragraph(), "9.1 Overview & Scope", 2)
    add_styled_paragraph(doc, "RanchiKart is a modular e-commerce solution comprising 6 core backend modules, 13 relational database tables, 35+ REST API endpoints, and a multi-page React Single Page Application.")

    style_heading(doc.add_paragraph(), "9.2 Module Breakdown & Descriptions", 2)
    module_data = [
        ["1. Auth & Security Module", "Handles registration, login, Google OAuth 2.0, magic links, TOTP 2FA, and WebAuthn passkeys.", "authController.ts, passkeyService.ts"],
        ["2. Catalog Module", "Manages 20 ProductKinds, categories tree, product search, filtering, and custom dimension bounds.", "catalogController.ts, catalogService.ts"],
        ["3. Cart & Order Module", "Manages user cart, order creation, custom text/size storage, and order state lifecycle.", "orderController.ts, orderService.ts"],
        ["4. Payment Module", "Handles Razorpay order creation, payment verification, and webhook cryptographic validation.", "paymentController.ts, paymentService.ts"],
        ["5. Coupon & Promo Module", "Handles percent/fixed coupon validation, category scoping, usage limits, and discount deduction.", "couponController.ts, couponService.ts"],
        ["6. Admin Back-Office", "Provides dashboard metrics, product/order CRUD, user ban/role control, and AdminLog audit review.", "adminController.ts, adminService.ts"]
    ]
    create_styled_table(doc, ["Module Name", "Functional Scope & Responsibility", "Key Code Artifacts"], module_data, [1.8, 3.2, 1.5])

    style_heading(doc.add_paragraph(), "9.3 Tools & Technology Stack", 2)
    tech_data = [
        ["Runtime Environment", "Bun 1.2 (Fast JS/TS bundler & runtime) / Node.js 22 LTS", "v1.2 / v22"],
        ["Backend Framework", "Fastify v5 (High performance, low overhead web framework)", "v5.2.0"],
        ["ORM & Database", "Prisma ORM 6.9 with PostgreSQL 17 (NeonDB managed)", "v6.9.0"],
        ["Frontend Framework", "React 18.3 SPA with Vite 5.4 build system & Tailwind CSS", "v18.3.1"],
        ["State Management", "Zustand 5 (Lightweight client state for Auth & Cart)", "v5.0.0"],
        ["Validation & Schema", "Zod 3.24 (TypeScript-first schema validation)", "v3.24.0"],
        ["Cache & Session", "Redis 7 (Redis Cloud managed instance)", "v7.0"],
        ["Payment Gateway", "Razorpay Node SDK & Client Checkout JS", "v2.9.5"]
    ]
    create_styled_table(doc, ["Layer / Component", "Technology Selected", "Version"], tech_data, [2.0, 3.5, 1.0])

    style_heading(doc.add_paragraph(), "Hardware Requirements", 3)
    hw_data = [
        ["Development Workstation", "Dual/Quad-Core Intel/AMD CPU, 8 GB RAM, 10 GB Free Storage, OS: Linux/Windows/macOS"],
        ["Production Server (Cloud)", "1 vCPU, 512 MB RAM (Render Free Tier / Vercel Edge Serverless)"],
        ["End-User Client Device", "Any Smartphone, Tablet, or Laptop with modern browser (Chrome, Firefox, Safari, Edge)"]
    ]
    create_styled_table(doc, ["Environment", "Minimum Hardware & System Requirements"], hw_data, [2.2, 4.3])

    style_heading(doc.add_paragraph(), "9.4 System Design", 2)
    
    style_heading(doc.add_paragraph(), "9.4.1 Data Flow Diagram (DFD)", 3)
    add_styled_paragraph(doc, "Level 0 (Context DFD): Represents RanchiKart as a single central process interacting with three external entities: User/Customer, Admin, and External Payment Gateway (Razorpay).")
    add_styled_paragraph(doc, "Level 1 DFD: Decomposes the system into 5 primary processes: 1.0 Authenticate User, 2.0 Browse & Customise Catalog, 3.0 Cart & Order Processing, 4.0 Payment Gateway Verification, and 5.0 Admin Store Management.")

    style_heading(doc.add_paragraph(), "9.4.2 Entity-Relationship (ER) Diagram Description", 3)
    add_styled_paragraph(doc, "The ER Model contains 13 primary entities centered around User, Product, and Order. A User can place multiple Orders (1:N) and own multiple PasskeyCredentials (1:N). A Category has a recursive parent-child hierarchy (1:N). A Product belongs to a Category (N:1) and has multiple ProductVariants (1:N). An Order contains multiple OrderItems (1:N) which link to specific Products and ProductVariants.")

    style_heading(doc.add_paragraph(), "9.4.3 Database Design", 3)
    
    # Storage Space Estimation Table (Guideline Requirement 9.5.iii.a)
    style_heading(doc.add_paragraph(), "a. Database Storage Space Estimation Math (10,000 Users Scale)", 4)
    db_math_data = [
        ["User", "512 Bytes", "10,000", "5.12 MB", "Includes email, Argon2 hash, tokens"],
        ["Category", "256 Bytes", "100", "0.03 MB", "Product categories & parent links"],
        ["Product", "1,024 Bytes", "2,500", "2.56 MB", "Product specs, images, dimension limits"],
        ["ProductVariant", "256 Bytes", "5,000", "1.28 MB", "Variant attributes & price deltas"],
        ["Order", "384 Bytes", "50,000", "19.20 MB", "Order totals, status, address JSON"],
        ["OrderItem", "256 Bytes", "125,000", "32.00 MB", "Line items, custom dimensions & text"],
        ["Payment", "512 Bytes", "50,000", "25.60 MB", "Razorpay IDs, signatures, responses"],
        ["Review / Wishlist", "256 Bytes", "20,000", "5.12 MB", "Ratings, review body, wishlist links"],
        ["AdminLog / Passkey", "256 Bytes", "15,000", "3.84 MB", "Audit logs & WebAuthn credentials"],
        ["Indexes & Overhead", "—", "—", "25.00 MB", "B-Tree indexes on FKs & unique keys"],
        ["Total DB Storage", "—", "—", "119.75 MB", "Comfortably fits within 512 MB Free Tier"]
    ]
    create_styled_table(doc, ["Table Name", "Est. Bytes / Record", "Est. Record Count", "Est. Table Size", "Storage Calculation Context"], db_math_data, [1.5, 1.2, 1.2, 1.2, 1.4])

    style_heading(doc.add_paragraph(), "b. Database Tables, Constraints & Relationships", 4)
    tbl_detail_data = [
        ["User", "id (PK CUID), email (UK), passwordHash, role (Enum), coins, isBanned, totpSecret", "PK: id, UK: email. FK relations to Order, Review, Wishlist, PasskeyCredential."],
        ["PasskeyCredential", "id (PK), userId (FK), credentialID (UK), credentialPublicKey (Bytes), counter", "FK: userId → User.id ON DELETE CASCADE."],
        ["Category", "id (PK), slug (UK), name, kind (Enum), parentId (FK recursive)", "FK: parentId → Category.id."],
        ["Product", "id (PK), categoryId (FK), slug (UK), basePrice, stock, minWidthMm, maxWidthMm", "FK: categoryId → Category.id."],
        ["ProductVariant", "id (PK), productId (FK), sku (UK), priceDelta, attributes (JSON)", "FK: productId → Product.id ON DELETE CASCADE."],
        ["Order", "id (PK), userId (FK), status (Enum), subtotal, shippingFee, total, address (JSON)", "FK: userId → User.id, FK: couponId → Coupon.id."],
        ["OrderItem", "id (PK), orderId (FK), productId (FK), customWidthMm, customHeightMm, customText", "FK: orderId → Order.id ON DELETE CASCADE."],
        ["Payment", "id (PK), orderId (FK), provider (Enum), providerOrderId (UK), providerSignature", "FK: orderId → Order.id ON DELETE CASCADE."],
        ["AdminLog", "id (PK), adminId (FK), action, entity, entityId, meta (JSON)", "FK: adminId → User.id."]
    ]
    create_styled_table(doc, ["Table Name", "Key Columns & Data Types", "Constraints & Foreign Keys"], tbl_detail_data, [1.5, 2.5, 2.5])

    style_heading(doc.add_paragraph(), "c. Details of Stored Procedures & Transactions", 4)
    add_styled_paragraph(doc, "RanchiKart executes atomic multi-table mutations via Prisma's $transaction API (which maps directly to BEGIN ... COMMIT blocks in PostgreSQL):")
    add_styled_bullet(doc, "Executes Order creation, OrderItem insertions, stock decrement, and Coupon usedCount increment in a single atomic transaction.", "Order Checkout Transaction: ")
    add_styled_bullet(doc, "Updates Payment status to CAPTURED, Order status to PAID, and logs administrative audit metadata in a single transaction block.", "Payment Reconciliation Transaction: ")

    style_heading(doc.add_paragraph(), "d. Database Triggers & Auto-Constraints", 4)
    add_styled_bullet(doc, "PostgreSQL default expression ((random() * 100)::int) automatically populates mock stock levels on product creation.", "Automatic Stock Generation: ")
    add_styled_bullet(doc, "Prisma @updatedAt decorator generates automated database triggers updating the updatedAt timestamp column upon any row modification.", "Automated Timestamps: ")
    add_styled_bullet(doc, "Foreign keys specify ON DELETE CASCADE for dependent rows (e.g. OrderItem, Payment, Review when parent is removed).", "Cascading Deletions: ")

    style_heading(doc.add_paragraph(), "e. Database Users, Roles & Permissions", 4)
    db_role_data = [
        ["postgres (DB Superuser)", "Full administrative control, schema migration execution, database creation", "Direct console / migration scripts"],
        ["neondb_owner (App User)", "CRUD operations on all 13 application tables (SELECT, INSERT, UPDATE, DELETE)", "Fastify Backend Connection Pool"],
        ["USER (App Role)", "Read active products/categories, create/read own cart, orders, reviews, addresses", "Application JWT Guard"],
        ["ADMIN (App Role)", "Full CRUD on products, categories, coupons, orders, users, and review of AdminLogs", "Admin Dashboard JWT Guard"]
    ]
    create_styled_table(doc, ["Role Name", "Granted Permissions & Access Rights", "Usage Scope"], db_role_data, [1.8, 3.2, 1.5])

    style_heading(doc.add_paragraph(), "9.4.4 Form Design", 3)
    
    style_heading(doc.add_paragraph(), "a. Layout & Specifications for Key Forms", 4)
    add_styled_bullet(doc, "Combined tabbed form supporting Password Login, Registration, Google OAuth button, Magic Link request, and TOTP 2FA prompt.", "1. AuthForm (AuthPage): ")
    add_styled_bullet(doc, "Interactive product view containing variant select dropdowns, custom width (mm), custom height (mm), and custom text area inputs.", "2. Custom Product Form: ")
    add_styled_bullet(doc, "Multi-step checkout form capturing shipping address, pincode, coupon code application, and payment gateway launcher.", "3. Checkout Form: ")
    add_styled_bullet(doc, "Admin management forms for product creation, category setup, coupon configuration, and order tracking update.", "4. Admin Back-Office Forms: ")

    style_heading(doc.add_paragraph(), "b. Validation Checks (Zod Client & Server Rules)", 4)
    val_data = [
        ["Email Field", "email().min(5).max(255)", "Must be a valid email format, non-empty, auto-trimmed."],
        ["Password Field", "string().min(8).regex(/^(?=.*[a-z])(?=.*[A-Z])(?=.*\\d)/)", "Minimum 8 chars, must contain uppercase, lowercase, and digit."],
        ["Phone Number", "string().regex(/^[6-9]\\d{9}$/)", "Must be a valid 10-digit Indian mobile number."],
        ["Pincode Field", "string().regex(/^\\d{6}$/)", "Must be a valid 6-digit postal code."],
        ["Custom Dimensions", "number().min(product.minWidthMm).max(product.maxWidthMm)", "Width and height must strictly fall within product boundaries."]
    ]
    create_styled_table(doc, ["Field Name", "Zod Validation Schema Rule", "Validation Behavior & Message"], val_data, [1.5, 2.5, 2.5])

    style_heading(doc.add_paragraph(), "c. Form Calculations", 4)
    add_styled_paragraph(doc, "1. Item Line Subtotal = (Product BasePrice + Variant PriceDelta) × Quantity")
    add_styled_paragraph(doc, "2. Custom Area Multiplier (if applicable) = (CustomWidthMm × CustomHeightMm) / (DefaultWidthMm × DefaultHeightMm)")
    add_styled_paragraph(doc, "3. Coupon Discount = If PERCENT: Subtotal × (Value / 100); If FIXED: Value in Paisa")
    add_styled_paragraph(doc, "4. Grand Total = Subtotal - DiscountAmount + ShippingFee")

    style_heading(doc.add_paragraph(), "d. Data Retrieval & Database Update Pipeline", 4)
    add_styled_paragraph(doc, "Form Submission → Axios Async Request (with Authorization: Bearer <JWT>) → Fastify Route Controller → Zod Schema Parse → Service Business Logic → Prisma ORM Query → PostgreSQL Database → JSON HTTP Response → Zustand State Hydration → React UI Re-render.")

    style_heading(doc.add_paragraph(), "9.4.5 Data Processing (Periodic Background Jobs)", 3)
    add_styled_bullet(doc, "Automated Redis TTL expiration purges stale JWT refresh tokens and expired rate-limiting counters.", "1. Token Cleanup: ")
    add_styled_bullet(doc, "Unverified email OTPs automatically expire after 10 minutes (emailOtpExpiry timestamp comparison).", "2. OTP Expiry: ")
    add_styled_bullet(doc, "Scheduled cleanup job permanently purges soft-deleted users (isDeleted = true) after 30 days retention.", "3. User Hard-Purge: ")

    style_heading(doc.add_paragraph(), "9.4.6 Networking Details & Protocols", 3)
    net_data = [
        ["Application Protocol", "HTTPS (TLS 1.3 encrypted REST API over JSON)"],
        ["CORS Security", "@fastify/cors configured with strict white-listed CORS_ORIGIN"],
        ["Security Headers", "@fastify/helmet providing Content-Security-Policy (CSP) & HSTS"],
        ["Rate Limiting", "@fastify/rate-limit enforcing maximum 100 requests / min per IP via Redis"]
    ]
    create_styled_table(doc, ["Networking Aspect", "Implementation Specification"], net_data, [2.0, 4.5])

    style_heading(doc.add_paragraph(), "9.4.7 Security Features & Vulnerability Mitigation", 3)
    sec_data = [
        ["Password Security", "Argon2id cryptographic hashing (memory-hard, salt-protected against rainbow tables)."],
        ["Authentication Session", "Dual JWT token model (15-min Access Token in memory, 7-day Refresh Token in HTTP-only cookie)."],
        ["Multi-Factor Auth (2FA)", "TOTP time-based 2FA via Google Authenticator + WebAuthn FIDO2 biometric passkeys."],
        ["OWASP Top-10 Defense", "SQL Injection prevented by Prisma parameterized queries; XSS prevented by React escaping & CSP."]
    ]
    create_styled_table(doc, ["Security Vector", "Technical Defense Mechanism"], sec_data, [2.0, 4.5])

    # ---------------------------------------------------------
    # 10.0 SYSTEM TESTING
    # ---------------------------------------------------------
    style_heading(doc.add_paragraph(), "10.0 SYSTEM TESTING", 1)
    
    style_heading(doc.add_paragraph(), "10.1 Testing Strategy", 2)
    add_styled_paragraph(doc, "RanchiKart employs an automated, layered testing strategy using Bun's built-in test runner alongside fastify.inject() (Fastify's in-process HTTP simulator) against a mocked Prisma client (prismaMockInit.ts). This allows complete request/response integration tests without requiring a live database connection during CI.")

    style_heading(doc.add_paragraph(), "10.2 Details of Each Test Case Used", 2)
    tc_data = [
        ["TC-01", "POST /auth/register — unique email", "201 Created; user object returned; single create() call", "Pass"],
        ["TC-02", "POST /auth/register — duplicate email", "409 Conflict; 'User with email already exists'", "Pass"],
        ["TC-03", "POST /auth/login — valid credentials", "200 OK; JWT access + refresh token pair issued", "Pass"],
        ["TC-04", "GET /products — catalog listing", "200 OK; paginated product list returned", "Pass"],
        ["TC-05", "GET /products?search=&category=", "200 OK; filtered/sorted results matching query params", "Pass"],
        ["TC-06", "Coupon apply — expired/invalid code", "4xx rejection with descriptive error, order unaffected", "Pass"],
        ["TC-07", "Order creation — insufficient stock", "4xx rejection; transaction aborted, order not created", "Pass"],
        ["TC-08", "Payment verification — bad signature", "Rejected; status remains PENDING_PAYMENT", "Pass"]
    ]
    create_styled_table(doc, ["Test Case", "Scenario / Test Vector", "Expected Output / Behavior", "Status"], tc_data, [1.0, 2.2, 2.5, 0.8])

    style_heading(doc.add_paragraph(), "10.3 Testing Report & Quality Metrics", 2)
    add_styled_paragraph(doc, "Continuous Integration (.github/workflows/CI.yaml) executes automatically on every push and pull request to the main branch. The build pipeline runs TypeScript type-checking (tsc --noEmit), unit/integration test suites, and production bundle builds for both backend and frontend. The repository currently maintains a 100% CI pass rate across all automated test suites.")

    # ---------------------------------------------------------
    # 11.0 & 12.0 SCOPE & FUTURE IMPLEMENTATION
    # ---------------------------------------------------------
    style_heading(doc.add_paragraph(), "11.0 SCOPE OF OUR SYSTEM", 1)
    add_styled_paragraph(doc, "In Scope: Single-storefront catalog management, 5-method authentication, Razorpay payment gateway integration, custom order item dimensions/text, coupon system, order tracking, product reviews/wishlist, and complete admin back-office.")
    add_styled_paragraph(doc, "Out of Scope (Current Release): Multi-vendor seller onboarding payouts, native mobile apps (iOS/Android), automated SMS/WhatsApp notifications, and third-party courier API automated tracking dispatch.")

    style_heading(doc.add_paragraph(), "12.0 FUTURE IMPLEMENTATION & ROADMAP", 1)
    add_styled_bullet(doc, "Onboarding independent local sellers with vendor-scoped catalogs and automated payout splits.", "1. Multi-Vendor Marketplace Mode: ")
    add_styled_bullet(doc, "Developing dedicated React Native apps or a Progressive Web App (PWA) with offline support.", "2. Mobile Applications: ")
    add_styled_bullet(doc, "Integrating Twilio / Fast2SMS and WhatsApp Business API for instant SMS updates.", "3. SMS & WhatsApp Notifications: ")
    add_styled_bullet(doc, "Direct integration with Shiprocket / Delhivery APIs for automatic shipping label generation.", "4. Logistics Partner Integration: ")
    add_styled_bullet(doc, "Vector embedding-based recommendations using pgvector for personalized search ranking.", "5. AI-Assisted Recommendations: ")
    add_styled_bullet(doc, "Automatic GST tax calculation and downloadable B2B tax invoices.", "6. GST-Compliant Invoicing: ")
    add_styled_bullet(doc, "Activating wallet cashback building on the existing coins column on the User model.", "7. Wallet & Cashback System: ")
    add_styled_bullet(doc, "Cohort analysis, sales forecasting charts, and automated low-stock alerts.", "8. Richer Admin Analytics: ")
    add_styled_bullet(doc, "Scheduled Redis cron jobs for abandoned cart recovery notifications.", "9. Abandoned Cart Recovery: ")
    add_styled_bullet(doc, "Third-party penetration testing and formal OWASP compliance verification.", "10. Security Audit: ")
    add_styled_bullet(doc, "Adding Hindi and Nagpuri language support for regional accessibility across Ranchi.", "11. Multilingual Support: ")

    # ---------------------------------------------------------
    # 13.0 BIBLIOGRAPHY
    # ---------------------------------------------------------
    style_heading(doc.add_paragraph(), "13.0 BIBLIOGRAPHY & REFERENCES", 1)
    refs = [
        "Fastify Web Framework — Official Documentation, https://fastify.dev",
        "React Library — Official Documentation, https://react.dev",
        "Prisma ORM — Technical Documentation & Guides, https://www.prisma.io/docs",
        "PostgreSQL 17 — Official Manual & Administration Guide, https://www.postgresql.org/docs",
        "Razorpay API Reference & Integration Guide, https://razorpay.com/docs",
        "W3C Web Authentication Spec (WebAuthn / FIDO2), https://www.w3.org/TR/webauthn-3/",
        "Zod Schema Validation — TypeScript Documentation, https://zod.dev",
        "Redis Cloud & In-Memory Data Store Documentation, https://redis.io/docs",
        "The ICFAI University Jharkhand, Ranchi — Guideline for SIP Project Report Writing (Faculty Guide: Prof. Abhay Kumar Sinha)",
        "RanchiKart GitHub Project Repository, Source Code & Database Architecture Schema"
    ]
    for r in refs:
        add_styled_bullet(doc, r)

    # ---------------------------------------------------------
    # APPENDICES
    # ---------------------------------------------------------
    doc.add_page_break()
    style_heading(doc.add_paragraph(), "APPENDIX A — SYSTEM FLOWCHART", 1)
    add_styled_paragraph(doc, "The flowchart below illustrates the end-to-end customer journey through RanchiKart, from initial page access through authentication, catalog navigation, custom dimension selection, cart checkout, Razorpay payment processing, and post-order tracking.")
    add_styled_paragraph(doc, "[User Access Site] → [Browse Catalog / Search] → [Select Product] → [Input Custom Width/Height/Text (if applicable)] → [Add to Cart] → [Authenticate / Register] → [Checkout Address & Apply Coupon] → [Launch Razorpay Gateway] → [HMAC Verification] → [Order Confirmed (PAID)] → [Admin Dispatch & Track].")

    doc.add_page_break()
    style_heading(doc.add_paragraph(), "APPENDIX B — PROJECT RESPONSIBILITY DOCUMENT (FULL)", 1)
    add_styled_paragraph(doc, "Project Topic: RanchiKart (BasketByte) — Fully Functional E-Commerce Portal for Ranchi.")
    add_styled_paragraph(doc, "Project Description: A full-stack e-commerce platform developed for the Ranchi market, focused on providing a secure, responsive, online shopping experience with efficient product management, user authentication, order management, and modern deployment practices.")
    
    resp_full_data = [
        ["1. Nishan", "Backend Development; DevOps & Deployment; Database Architecture; Security"],
        ["2. Aditya", "Requirement Analysis; Research; UI/UX Design; Frontend Development; State Management; Optimization"],
        ["3. Suraj", "Quality Assurance & Testing Lead; Functional, UI, Validation, Smoke, Regression, Compatibility, E2E Testing"],
        ["4. Tushar", "Requirement Analysis; System Design (DFD/ER); Technical Documentation; Monitoring"],
        ["5. Adarsh", "Project Coordination; Product Management; Content Management; UAT Testing; Presentation"]
    ]
    create_styled_table(doc, ["Team Member", "Detailed Responsibility & Area Breakdown"], resp_full_data, [1.8, 4.7])

    # Output path
    output_path = "/home/nishu/TechStack/codes/RanchiKart/RanchiKart_SIP_Project_Report_Formatted.docx"
    doc.save(output_path)
    print(f"Report successfully built and saved to: {output_path}")

if __name__ == "__main__":
    build_sip_report()
